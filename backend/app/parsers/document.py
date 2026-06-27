from pathlib import Path

from backend.app.domain.enums import CostLevel, FileType, LatencyLevel, Modality
from backend.app.parsers.base import BaseParser, ParseRequest, ParseResult, ParserMetadata


class PdfNativeParser(BaseParser):
    metadata = ParserMetadata(
        parser_id="pdf_native_text",
        name="PDF Native Text Parser",
        supported_file_types=[FileType.PDF],
        supported_modalities=[Modality.DOCUMENT, Modality.TEXT],
        cost_level=CostLevel.LOW,
        latency_level=LatencyLevel.LOW,
        expected_quality=0.78,
        version="0.1.0",
        enabled=True,
    )

    def parse(self, request: ParseRequest) -> ParseResult:
        if request.storage_path:
            path = Path(request.storage_path)
            try:
                import fitz  # type: ignore[import-not-found]

                pages: list[str] = []
                layout_blocks: list[dict[str, object]] = []
                with fitz.open(path) as document:
                    for page_index, page in enumerate(document):
                        page_text = page.get_text("text").strip()
                        if page_text:
                            pages.append(page_text)
                        for block_index, block in enumerate(page.get_text("blocks")):
                            if len(block) < 5:
                                continue
                            block_text = str(block[4]).strip()
                            if not block_text:
                                continue
                            layout_blocks.append(
                                {
                                    "page": page_index + 1,
                                    "block_id": f"page-{page_index + 1}-block-{block_index}",
                                    "type": "text",
                                    "bbox": [block[0], block[1], block[2], block[3]],
                                    "text": block_text,
                                }
                            )

                parsed_text = "\n\n".join(pages).strip()
                return ParseResult(
                    parser_id=self.parser_id,
                    parsed_text=parsed_text,
                    layout_blocks=layout_blocks,
                    structured_data={
                        "source": "pdf_native_text",
                        "file_id": request.file_id,
                        "page_count": len(pages),
                    },
                    confidence_score=0.86 if parsed_text else 0.25,
                    warnings=[] if parsed_text else ["No native PDF text layer was detected."],
                )
            except ImportError:
                preview = self._read_text_preview(request)
                return ParseResult(
                    parser_id=self.parser_id,
                    parsed_text=(
                        preview or f"Mock PDF native text extracted from {request.filename}."
                    ),
                    structured_data={"source": "pdf_native_text", "file_id": request.file_id},
                    confidence_score=0.56 if preview else 0.35,
                    warnings=["PyMuPDF is not installed; real PDF extraction was skipped."],
                )
            except Exception as exc:
                return ParseResult(
                    parser_id=self.parser_id,
                    parsed_text=None,
                    structured_data={"source": "pdf_native_text", "file_id": request.file_id},
                    confidence_score=0.0,
                    warnings=[f"PDF extraction failed: {exc}"],
                )

        preview = self._read_text_preview(request)
        return ParseResult(
            parser_id=self.parser_id,
            parsed_text=preview or f"Mock PDF native text extracted from {request.filename}.",
            structured_data={"source": "pdf_native_text", "file_id": request.file_id},
            confidence_score=0.56 if preview else 0.35,
            warnings=["No storage path was provided for PDF extraction."],
        )


class DocxParser(BaseParser):
    metadata = ParserMetadata(
        parser_id="docx_text",
        name="DOCX Parser",
        supported_file_types=[FileType.DOCX],
        supported_modalities=[Modality.DOCUMENT, Modality.TEXT],
        cost_level=CostLevel.LOW,
        latency_level=LatencyLevel.LOW,
        expected_quality=0.76,
        version="0.1.0",
        enabled=True,
    )

    def parse(self, request: ParseRequest) -> ParseResult:
        if request.storage_path:
            path = Path(request.storage_path)
            try:
                from docx import Document  # type: ignore[import-not-found]

                document = Document(path)
                paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
                paragraphs = [paragraph for paragraph in paragraphs if paragraph]
                tables: list[dict[str, object]] = []
                for table_index, table in enumerate(document.tables):
                    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                    tables.append(
                        {
                            "table_id": f"table-{table_index}",
                            "row_count": len(rows),
                            "column_count": max((len(row) for row in rows), default=0),
                            "rows": rows,
                        }
                    )

                parsed_text = "\n\n".join(paragraphs)
                return ParseResult(
                    parser_id=self.parser_id,
                    parsed_text=parsed_text,
                    tables=tables,
                    layout_blocks=[
                        {"type": "paragraph", "index": index, "text": paragraph}
                        for index, paragraph in enumerate(paragraphs)
                    ],
                    structured_data={
                        "source": "docx_text",
                        "file_id": request.file_id,
                        "paragraph_count": len(paragraphs),
                        "table_count": len(tables),
                    },
                    confidence_score=0.88 if parsed_text or tables else 0.3,
                    warnings=(
                        [] if parsed_text or tables else ["DOCX contained no extractable text."]
                    ),
                )
            except ImportError:
                preview = self._read_text_preview(request)
                return ParseResult(
                    parser_id=self.parser_id,
                    parsed_text=preview or f"Mock DOCX text extracted from {request.filename}.",
                    structured_data={"source": "docx_text", "file_id": request.file_id},
                    confidence_score=0.55 if preview else 0.35,
                    warnings=["python-docx is not installed; real DOCX extraction was skipped."],
                )
            except Exception as exc:
                return ParseResult(
                    parser_id=self.parser_id,
                    parsed_text=None,
                    structured_data={"source": "docx_text", "file_id": request.file_id},
                    confidence_score=0.0,
                    warnings=[f"DOCX extraction failed: {exc}"],
                )

        preview = self._read_text_preview(request)
        return ParseResult(
            parser_id=self.parser_id,
            parsed_text=preview or f"Mock DOCX text extracted from {request.filename}.",
            structured_data={"source": "docx_text", "file_id": request.file_id},
            confidence_score=0.55 if preview else 0.35,
            warnings=["No storage path was provided for DOCX extraction."],
        )


class HtmlParser(BaseParser):
    metadata = ParserMetadata(
        parser_id="html_text",
        name="HTML Parser",
        supported_file_types=[FileType.HTML],
        supported_modalities=[Modality.DOCUMENT, Modality.TEXT],
        cost_level=CostLevel.LOW,
        latency_level=LatencyLevel.LOW,
        expected_quality=0.74,
        version="0.1.0",
        enabled=True,
    )

    def parse(self, request: ParseRequest) -> ParseResult:
        preview = self._read_text(request)
        tables: list[dict[str, object]] = []
        image_descriptions: list[dict[str, object]] = []

        try:
            from bs4 import BeautifulSoup  # type: ignore[import-not-found]

            soup = BeautifulSoup(preview, "html.parser")
            for element in soup(["script", "style", "noscript"]):
                element.decompose()
            text = soup.get_text("\n", strip=True)
            for table_index, table in enumerate(soup.find_all("table")):
                rows = []
                for row in table.find_all("tr"):
                    cells = [cell.get_text(" ", strip=True) for cell in row.find_all(["th", "td"])]
                    if cells:
                        rows.append(cells)
                tables.append(
                    {
                        "table_id": f"table-{table_index}",
                        "row_count": len(rows),
                        "column_count": max((len(row) for row in rows), default=0),
                        "rows": rows,
                    }
                )
            for image_index, image in enumerate(soup.find_all("img")):
                image_descriptions.append(
                    {
                        "image_id": f"image-{image_index}",
                        "src": image.get("src"),
                        "alt": image.get("alt"),
                    }
                )
            stripped = text
        except ImportError:
            stripped = preview.replace("<", " <").replace(">", "> ")

        return ParseResult(
            parser_id=self.parser_id,
            parsed_text=stripped or f"Mock HTML text extracted from {request.filename}.",
            tables=tables,
            image_descriptions=image_descriptions,
            layout_blocks=[
                {"type": "text", "index": index, "text": line}
                for index, line in enumerate(stripped.splitlines())
                if line.strip()
            ],
            structured_data={
                "source": "html_text",
                "file_id": request.file_id,
                "table_count": len(tables),
                "image_count": len(image_descriptions),
            },
            confidence_score=0.82 if stripped else 0.35,
            warnings=[] if stripped else ["HTML contained no extractable text."],
        )
